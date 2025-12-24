"""Read file tool with support for various file formats."""

from pathlib import Path
from typing import Optional

from opencode.tools.base import Tool, ToolResult


# Maximum lines to display to user (full content still sent to LLM)
MAX_DISPLAY_LINES = 10

# Threshold for "large file" handling
LARGE_FILE_THRESHOLD = 500

# Lines to show as preview for large files
LARGE_FILE_PREVIEW_LINES = 50

# File type categories
OFFICE_EXTENSIONS = {
    # Excel
    ".xlsx", ".xls", ".xlsm", ".xlsb",
    # Word
    ".docx", ".doc",
    # CSV (handled specially)
    ".csv",
}

# XML/structured data extensions
XML_EXTENSIONS = {".xml", ".xhtml", ".svg", ".plist", ".rss", ".atom"}

# Check for optional dependencies
try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import xlrd
    HAS_XLRD = True
except ImportError:
    HAS_XLRD = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    try:
        import PyPDF2 as pypdf
        HAS_PYPDF = True
    except ImportError:
        HAS_PYPDF = False


def read_excel_file(file_path: Path, sheet: Optional[str] = None) -> tuple[str, str]:
    """Read Excel file and return content as text.

    Args:
        file_path: Path to Excel file.
        sheet: Optional sheet name (default: first sheet).

    Returns:
        Tuple of (content, format_description)
    """
    suffix = file_path.suffix.lower()

    # Modern Excel (.xlsx, .xlsm, .xlsb)
    if suffix in (".xlsx", ".xlsm", ".xlsb"):
        if not HAS_OPENPYXL:
            return "", "Excel support requires: pip install openpyxl"

        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        sheet_names = wb.sheetnames

        if sheet and sheet in sheet_names:
            ws = wb[sheet]
        else:
            ws = wb.active

        lines = []
        lines.append(f"[EXCEL FILE: {file_path.name}]")
        lines.append(f"Sheets: {', '.join(sheet_names)}")
        lines.append(f"Active sheet: {ws.title}")
        lines.append("")

        # Read all rows
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            row_count += 1
            # Convert cells to strings, handle None
            cells = [str(cell) if cell is not None else "" for cell in row]
            lines.append(f"{row_count:4d} | " + " | ".join(cells))

        wb.close()
        return "\n".join(lines), f"Excel ({row_count} rows)"

    # Legacy Excel (.xls)
    elif suffix == ".xls":
        if not HAS_XLRD:
            return "", "Legacy Excel (.xls) support requires: pip install xlrd"

        wb = xlrd.open_workbook(str(file_path))
        sheet_names = wb.sheet_names()

        if sheet and sheet in sheet_names:
            ws = wb.sheet_by_name(sheet)
        else:
            ws = wb.sheet_by_index(0)

        lines = []
        lines.append(f"[EXCEL FILE: {file_path.name}]")
        lines.append(f"Sheets: {', '.join(sheet_names)}")
        lines.append(f"Active sheet: {ws.name}")
        lines.append("")

        for row_idx in range(ws.nrows):
            row = ws.row_values(row_idx)
            cells = [str(cell) if cell else "" for cell in row]
            lines.append(f"{row_idx + 1:4d} | " + " | ".join(cells))

        return "\n".join(lines), f"Excel ({ws.nrows} rows)"

    return "", f"Unsupported Excel format: {suffix}"


def read_word_file(file_path: Path) -> tuple[str, str]:
    """Read Word document and return content as text.

    Args:
        file_path: Path to Word document.

    Returns:
        Tuple of (content, format_description)
    """
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        if not HAS_DOCX:
            return "", "Word support requires: pip install python-docx"

        doc = DocxDocument(file_path)

        lines = []
        lines.append(f"[WORD DOCUMENT: {file_path.name}]")

        # Document properties
        if doc.core_properties.title:
            lines.append(f"Title: {doc.core_properties.title}")
        if doc.core_properties.author:
            lines.append(f"Author: {doc.core_properties.author}")
        lines.append("")

        # Read paragraphs
        para_count = 0
        for para in doc.paragraphs:
            if para.text.strip():
                para_count += 1
                # Check if it's a heading
                if para.style and para.style.name.startswith('Heading'):
                    lines.append(f"\n## {para.text}")
                else:
                    lines.append(para.text)

        # Read tables
        if doc.tables:
            lines.append(f"\n[TABLES: {len(doc.tables)}]")
            for table_idx, table in enumerate(doc.tables):
                lines.append(f"\n--- Table {table_idx + 1} ---")
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    lines.append(" | ".join(cells))

        return "\n".join(lines), f"Word ({para_count} paragraphs)"

    elif suffix == ".doc":
        # Legacy .doc format - try antiword or textract if available
        try:
            import subprocess
            result = subprocess.run(
                ["antiword", str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                content = f"[WORD DOCUMENT: {file_path.name}]\n\n{result.stdout}"
                return content, "Word (legacy .doc)"
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        return "", "Legacy Word (.doc) requires 'antiword' command or convert to .docx"

    return "", f"Unsupported Word format: {suffix}"


def read_xml_file(file_path: Path) -> tuple[str, str]:
    """Read XML file and return formatted content.

    Args:
        file_path: Path to XML file.

    Returns:
        Tuple of (content, format_description)
    """
    try:
        import xml.etree.ElementTree as ET
        from xml.dom import minidom

        # Try to parse and pretty-print
        tree = ET.parse(file_path)
        root = tree.getroot()

        # Convert to pretty-printed string
        xml_str = ET.tostring(root, encoding='unicode')
        try:
            # Try to format nicely
            dom = minidom.parseString(xml_str)
            pretty_xml = dom.toprettyxml(indent="  ")
            # Remove extra blank lines
            lines = [line for line in pretty_xml.splitlines() if line.strip()]
            content = "\n".join(lines)
        except Exception:
            content = xml_str

        header = f"[XML FILE: {file_path.name}]\n"
        header += f"Root element: <{root.tag}>\n"
        if root.attrib:
            header += f"Attributes: {root.attrib}\n"
        header += f"Child elements: {len(list(root))}\n\n"

        return header + content, f"XML ({len(list(root))} elements)"

    except ET.ParseError as e:
        # If parsing fails, read as plain text
        content = file_path.read_text(encoding='utf-8', errors='replace')
        return f"[XML FILE: {file_path.name}]\n[Parse warning: {e}]\n\n{content}", "XML (raw)"
    except Exception as e:
        return "", f"Failed to read XML: {e}"


def read_csv_file(file_path: Path) -> tuple[str, str]:
    """Read CSV file and return content as text.

    Args:
        file_path: Path to CSV file.

    Returns:
        Tuple of (content, format_description)
    """
    import csv

    lines = []
    lines.append(f"[CSV FILE: {file_path.name}]")
    lines.append("")

    row_count = 0
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        reader = csv.reader(f)
        for row in reader:
            row_count += 1
            lines.append(f"{row_count:4d} | " + " | ".join(row))

    return "\n".join(lines), f"CSV ({row_count} rows)"


def read_pdf_file(file_path: Path) -> tuple[str, str]:
    """Read PDF file and return content as text.

    Args:
        file_path: Path to PDF file.

    Returns:
        Tuple of (content, format_description)
    """
    if not HAS_PYPDF:
        return "", "PDF support requires: pip install pypdf (or PyPDF2)"

    lines = []
    lines.append(f"[PDF FILE: {file_path.name}]")

    try:
        reader = pypdf.PdfReader(str(file_path))
        num_pages = len(reader.pages)
        lines.append(f"Pages: {num_pages}")

        # Document metadata
        if reader.metadata:
            if reader.metadata.title:
                lines.append(f"Title: {reader.metadata.title}")
            if reader.metadata.author:
                lines.append(f"Author: {reader.metadata.author}")

        lines.append("")

        # Extract text from each page
        total_chars = 0
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text and text.strip():
                lines.append(f"--- Page {page_num} ---")
                lines.append(text.strip())
                lines.append("")
                total_chars += len(text)

        if total_chars == 0:
            lines.append("[Note: No extractable text found. This PDF may contain only images.]")

        return "\n".join(lines), f"PDF ({num_pages} pages)"

    except Exception as e:
        return "", f"Failed to read PDF: {e}"


class ReadTool(Tool):
    """Read file contents with support for various formats."""

    name = "read"
    description = (
        "Read the contents of a file. Supports: "
        "text files, code files, Excel (.xlsx, .xls), Word (.docx), PDF, CSV, XML, and more. "
        "For large files (>500 lines), returns outline + preview by default. "
        "Use full=true to read entire file. Use sheet parameter for Excel files."
    )
    requires_build_mode = False

    def execute(
        self,
        path: str,
        lines: str = None,
        full: bool = False,
        sheet: str = None
    ) -> ToolResult:
        """Read a file's contents.

        Args:
            path: Path to the file.
            lines: Optional line range (e.g., "10-20" or "10") for text files.
            full: If True, read entire file even if large.
            sheet: Sheet name for Excel files (default: first sheet).

        Returns:
            ToolResult with file contents.
        """
        self._check_mode()

        try:
            file_path = self._resolve_path(path)
        except ValueError as e:
            return ToolResult.fail(str(e))

        if not file_path.exists():
            return ToolResult.fail(f"File not found: {path}")

        if not file_path.is_file():
            return ToolResult.fail(f"Not a file: {path}")

        # Check for special file types
        suffix = file_path.suffix.lower()

        # Handle Excel files
        if suffix in (".xlsx", ".xls", ".xlsm", ".xlsb"):
            return self._read_excel(file_path, sheet)

        # Handle Word files
        if suffix in (".docx", ".doc"):
            return self._read_word(file_path)

        # Handle PDF files
        if suffix == ".pdf":
            return self._read_pdf(file_path)

        # Handle CSV files
        if suffix == ".csv":
            return self._read_csv(file_path)

        # Handle XML files
        if suffix in XML_EXTENSIONS:
            return self._read_xml(file_path)

        # Handle regular text files
        try:
            content = file_path.read_text(encoding='utf-8')
            file_lines = content.splitlines()
            total_lines = len(file_lines)
        except UnicodeDecodeError:
            # Try binary detection for unknown files
            return self._handle_binary_file(file_path)

        if lines:
            # Specific line range requested - return those lines
            start, end = self._parse_line_range(lines, total_lines)
            selected_lines = file_lines[start:end]
            numbered = [
                f"{i + start + 1:4d} | {line}"
                for i, line in enumerate(selected_lines)
            ]
            full_output = "\n".join(numbered)
            display_lines = numbered

        elif total_lines > LARGE_FILE_THRESHOLD and not full:
            # Large file - provide outline + preview for LLM (unless full read requested)
            return self._handle_large_file(path, file_path, file_lines, total_lines)

        else:
            # Small file - return everything
            numbered = [
                f"{i + 1:4d} | {line}"
                for i, line in enumerate(file_lines)
            ]
            full_output = "\n".join(numbered)
            display_lines = numbered

        # Truncated display for user
        if len(display_lines) > MAX_DISPLAY_LINES:
            preview = "\n".join(display_lines[:MAX_DISPLAY_LINES])
            remaining = len(display_lines) - MAX_DISPLAY_LINES
            print(f"\033[34m> Reading {path} ({total_lines} lines)\033[0m")
            print(preview)
            print(f"\033[90m   ... ({remaining} more lines)\033[0m")
        else:
            print(f"\033[34m> Reading {path} ({total_lines} lines)\033[0m")
            print("\n".join(display_lines))

        # Return full content to LLM
        return ToolResult(success=True, output="", _llm_output=full_output)

    def _handle_large_file(
        self,
        path: str,
        file_path: Path,
        file_lines: list,
        total_lines: int
    ) -> ToolResult:
        """Handle large files by providing outline + preview.

        For files over LARGE_FILE_THRESHOLD lines, we provide:
        1. File structure/outline (if available)
        2. First N lines as preview
        3. Instructions for navigating to specific sections
        """
        # Try to get outline
        outline_text = self._get_outline(path, file_path)

        # Build preview of first N lines
        preview_lines = file_lines[:LARGE_FILE_PREVIEW_LINES]
        numbered_preview = [
            f"{i + 1:4d} | {line}"
            for i, line in enumerate(preview_lines)
        ]
        preview_text = "\n".join(numbered_preview)

        # Build LLM output
        llm_parts = [
            f"[LARGE FILE: {total_lines} lines]",
            f"Path: {path}",
            "",
        ]

        if outline_text:
            llm_parts.extend([
                "[STRUCTURE]",
                outline_text,
                "",
            ])

        llm_parts.extend([
            f"[PREVIEW: First {LARGE_FILE_PREVIEW_LINES} lines]",
            preview_text,
            "",
            f"[TIP: Use read(path=\"{path}\", lines=\"START-END\") to view specific sections]",
        ])

        llm_output = "\n".join(llm_parts)

        # Display output for user
        print(f"\033[34m> Reading {path} ({total_lines} lines - large file)\033[0m")
        if outline_text:
            # Show abbreviated outline
            outline_lines = outline_text.splitlines()
            if len(outline_lines) > 10:
                print("\n".join(outline_lines[:10]))
                print(f"\033[90m   ... ({len(outline_lines) - 10} more symbols)\033[0m")
            else:
                print(outline_text)
        print(f"\033[90m   [Showing structure + first {LARGE_FILE_PREVIEW_LINES} lines to LLM]\033[0m")

        return ToolResult(success=True, output="", _llm_output=llm_output)

    def _get_outline(self, path: str, file_path: Path) -> str:
        """Try to get file outline using OutlineTool patterns."""
        import re

        # Language detection
        LANGUAGE_MAP = {
            ".py": "python",
            ".js": "javascript",
            ".jsx": "javascript",
            ".ts": "typescript",
            ".tsx": "typescript",
            ".rs": "rust",
            ".go": "go",
            ".java": "java",
            ".c": "c",
            ".h": "c",
            ".cpp": "cpp",
            ".hpp": "cpp",
            ".rb": "ruby",
        }

        # Symbol patterns (simplified from OutlineTool)
        PATTERNS = {
            "python": [
                (r"^(\s*)class\s+(\w+)", "class"),
                (r"^(\s*)def\s+(\w+)", "function"),
            ],
            "javascript": [
                (r"^(\s*)class\s+(\w+)", "class"),
                (r"^(\s*)function\s+(\w+)", "function"),
                (r"^(\s*)const\s+(\w+)\s*=.*=>", "arrow"),
            ],
            "typescript": [
                (r"^(\s*)class\s+(\w+)", "class"),
                (r"^(\s*)interface\s+(\w+)", "interface"),
                (r"^(\s*)function\s+(\w+)", "function"),
                (r"^(\s*)type\s+(\w+)", "type"),
            ],
            "rust": [
                (r"^(\s*)(?:pub\s+)?struct\s+(\w+)", "struct"),
                (r"^(\s*)(?:pub\s+)?enum\s+(\w+)", "enum"),
                (r"^(\s*)(?:pub\s+)?fn\s+(\w+)", "fn"),
                (r"^(\s*)impl.*?(\w+)\s*\{", "impl"),
            ],
            "go": [
                (r"^type\s+(\w+)\s+struct", "struct"),
                (r"^func\s+(\w+)", "func"),
                (r"^func\s+\([^)]+\)\s+(\w+)", "method"),
            ],
        }

        suffix = file_path.suffix
        language = LANGUAGE_MAP.get(suffix)
        if not language or language not in PATTERNS:
            return ""

        patterns = [(re.compile(p), k) for p, k in PATTERNS.get(language, [])]
        if not patterns:
            return ""

        try:
            content = file_path.read_text(encoding='utf-8', errors='ignore')
        except Exception:
            return ""

        symbols = []
        for line_num, line in enumerate(content.splitlines(), 1):
            for pattern, kind in patterns:
                match = pattern.match(line)
                if match:
                    # Get name from last group
                    name = match.group(match.lastindex)
                    indent = len(match.group(1)) if match.lastindex > 1 else 0
                    symbols.append((indent, kind, name, line_num))
                    break

        if not symbols:
            return ""

        # Format as simple tree
        lines = []
        for indent, kind, name, line_num in symbols:
            prefix = "  " * (indent // 4)
            lines.append(f"{prefix}[{kind[0].upper()}] {name} ::{line_num}")

        return "\n".join(lines)

    def _parse_line_range(self, lines: str, total: int) -> tuple[int, int]:
        """Parse a line range like '10-20' or '10'."""
        if "-" in lines:
            parts = lines.split("-")
            start = int(parts[0]) - 1  # Convert to 0-indexed
            end = int(parts[1])
        else:
            start = int(lines) - 1
            end = start + 1
        return max(0, start), min(total, end)

    def _read_excel(self, file_path: Path, sheet: str = None) -> ToolResult:
        """Read Excel file."""
        try:
            content, description = read_excel_file(file_path, sheet)
            if not content:
                return ToolResult.fail(description)

            # Display preview
            lines = content.splitlines()
            print(f"\033[34m> Reading {file_path.name} ({description})\033[0m")
            if len(lines) > MAX_DISPLAY_LINES:
                print("\n".join(lines[:MAX_DISPLAY_LINES]))
                print(f"\033[90m   ... ({len(lines) - MAX_DISPLAY_LINES} more lines)\033[0m")
            else:
                print(content)

            return ToolResult(success=True, output="", _llm_output=content)
        except Exception as e:
            return ToolResult.fail(f"Failed to read Excel file: {e}")

    def _read_word(self, file_path: Path) -> ToolResult:
        """Read Word document."""
        try:
            content, description = read_word_file(file_path)
            if not content:
                return ToolResult.fail(description)

            # Display preview
            lines = content.splitlines()
            print(f"\033[34m> Reading {file_path.name} ({description})\033[0m")
            if len(lines) > MAX_DISPLAY_LINES:
                print("\n".join(lines[:MAX_DISPLAY_LINES]))
                print(f"\033[90m   ... ({len(lines) - MAX_DISPLAY_LINES} more lines)\033[0m")
            else:
                print(content)

            return ToolResult(success=True, output="", _llm_output=content)
        except Exception as e:
            return ToolResult.fail(f"Failed to read Word document: {e}")

    def _read_csv(self, file_path: Path) -> ToolResult:
        """Read CSV file."""
        try:
            content, description = read_csv_file(file_path)

            # Display preview
            lines = content.splitlines()
            print(f"\033[34m> Reading {file_path.name} ({description})\033[0m")
            if len(lines) > MAX_DISPLAY_LINES:
                print("\n".join(lines[:MAX_DISPLAY_LINES]))
                print(f"\033[90m   ... ({len(lines) - MAX_DISPLAY_LINES} more lines)\033[0m")
            else:
                print(content)

            return ToolResult(success=True, output="", _llm_output=content)
        except Exception as e:
            return ToolResult.fail(f"Failed to read CSV file: {e}")

    def _read_pdf(self, file_path: Path) -> ToolResult:
        """Read PDF file."""
        try:
            content, description = read_pdf_file(file_path)
            if not content:
                return ToolResult.fail(description)

            # Display preview
            lines = content.splitlines()
            print(f"\033[34m> Reading {file_path.name} ({description})\033[0m")
            if len(lines) > MAX_DISPLAY_LINES:
                print("\n".join(lines[:MAX_DISPLAY_LINES]))
                print(f"\033[90m   ... ({len(lines) - MAX_DISPLAY_LINES} more lines)\033[0m")
            else:
                print(content)

            return ToolResult(success=True, output="", _llm_output=content)
        except Exception as e:
            return ToolResult.fail(f"Failed to read PDF file: {e}")

    def _read_xml(self, file_path: Path) -> ToolResult:
        """Read XML file."""
        try:
            content, description = read_xml_file(file_path)
            if not content:
                return ToolResult.fail(description)

            # Display preview
            lines = content.splitlines()
            print(f"\033[34m> Reading {file_path.name} ({description})\033[0m")
            if len(lines) > MAX_DISPLAY_LINES:
                print("\n".join(lines[:MAX_DISPLAY_LINES]))
                print(f"\033[90m   ... ({len(lines) - MAX_DISPLAY_LINES} more lines)\033[0m")
            else:
                print(content)

            return ToolResult(success=True, output="", _llm_output=content)
        except Exception as e:
            return ToolResult.fail(f"Failed to read XML file: {e}")

    def _handle_binary_file(self, file_path: Path) -> ToolResult:
        """Handle binary files that can't be read as text."""
        suffix = file_path.suffix.lower()
        size = file_path.stat().st_size

        # Provide info about the file
        content = f"""[BINARY FILE: {file_path.name}]
Type: {suffix or 'unknown'}
Size: {size:,} bytes

This file cannot be read as text.
For specific file types, install optional dependencies:
  - Excel (.xlsx): pip install openpyxl
  - Excel (.xls): pip install xlrd
  - Word (.docx): pip install python-docx
  - PDF (.pdf): pip install pypdf
"""
        print(f"\033[33m> Binary file: {file_path.name} ({size:,} bytes)\033[0m")
        return ToolResult(success=True, output="", _llm_output=content)

    def get_schema(self) -> dict:
        """Return JSON schema for LLM function calling."""
        return {
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Path to the file to read. Supports: text, code, Excel (.xlsx/.xls), Word (.docx), PDF (.pdf), CSV (.csv), XML (.xml)"
                },
                "lines": {
                    "type": "string",
                    "description": "Optional line range for text files (e.g., '10-20' or '10')"
                },
                "full": {
                    "type": "boolean",
                    "description": "Read entire file even if large. Use when asked to understand, analyze, or explain code."
                },
                "sheet": {
                    "type": "string",
                    "description": "Sheet name for Excel files (default: first/active sheet)"
                }
            },
            "required": ["path"]
        }
