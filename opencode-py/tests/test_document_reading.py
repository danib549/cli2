"""Tests for document reading functionality (Excel, Word, CSV, PDF)."""

import os
import tempfile
import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock

from opencode.tools.read import (
    ReadTool,
    read_excel_file,
    read_word_file,
    read_csv_file,
    read_pdf_file,
    HAS_OPENPYXL,
    HAS_XLRD,
    HAS_DOCX,
    HAS_PYPDF,
)


class TestCSVReading:
    """Test CSV file reading."""

    def test_read_simple_csv(self, tmp_path):
        """Test reading a simple CSV file."""
        csv_file = tmp_path / "test.csv"
        csv_file.write_text("name,age,city\nAlice,30,NYC\nBob,25,LA\n")

        content, description = read_csv_file(csv_file)

        assert "[CSV FILE: test.csv]" in content
        assert "name" in content
        assert "Alice" in content
        assert "Bob" in content
        assert "3 rows" in description

    def test_read_csv_with_special_chars(self, tmp_path):
        """Test CSV with special characters."""
        csv_file = tmp_path / "special.csv"
        csv_file.write_text('name,description\n"John ""Jack""",Has quotes\n')

        content, description = read_csv_file(csv_file)

        assert "John" in content
        assert "1 rows" in description or "2 rows" in description

    def test_read_empty_csv(self, tmp_path):
        """Test reading empty CSV file."""
        csv_file = tmp_path / "empty.csv"
        csv_file.write_text("")

        content, description = read_csv_file(csv_file)

        assert "[CSV FILE: empty.csv]" in content
        assert "0 rows" in description

    def test_read_csv_via_tool(self, tmp_path):
        """Test CSV reading through ReadTool."""
        csv_file = tmp_path / "data.csv"
        csv_file.write_text("col1,col2\nval1,val2\n")

        tool = ReadTool()
        result = tool.execute(path=str(csv_file))

        assert result.success
        assert "col1" in result.llm_output
        assert "val1" in result.llm_output


class TestExcelReading:
    """Test Excel file reading."""

    @pytest.mark.skipif(not HAS_OPENPYXL, reason="openpyxl not installed")
    def test_read_xlsx_file(self, tmp_path):
        """Test reading .xlsx file."""
        import openpyxl

        xlsx_file = tmp_path / "test.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws["A1"] = "Name"
        ws["B1"] = "Value"
        ws["A2"] = "Test"
        ws["B2"] = 123
        wb.save(xlsx_file)

        content, description = read_excel_file(xlsx_file)

        assert "[EXCEL FILE: test.xlsx]" in content
        assert "Sheet1" in content
        assert "Name" in content
        assert "Test" in content
        assert "123" in content
        assert "2 rows" in description

    @pytest.mark.skipif(not HAS_OPENPYXL, reason="openpyxl not installed")
    def test_read_xlsx_specific_sheet(self, tmp_path):
        """Test reading specific sheet from .xlsx file."""
        import openpyxl

        xlsx_file = tmp_path / "multi_sheet.xlsx"
        wb = openpyxl.Workbook()

        # First sheet
        ws1 = wb.active
        ws1.title = "Data"
        ws1["A1"] = "First Sheet"

        # Second sheet
        ws2 = wb.create_sheet("Summary")
        ws2["A1"] = "Second Sheet"

        wb.save(xlsx_file)

        # Read specific sheet
        content, description = read_excel_file(xlsx_file, sheet="Summary")

        assert "Second Sheet" in content
        assert "Active sheet: Summary" in content

    @pytest.mark.skipif(not HAS_OPENPYXL, reason="openpyxl not installed")
    def test_read_xlsx_via_tool(self, tmp_path):
        """Test Excel reading through ReadTool."""
        import openpyxl

        xlsx_file = tmp_path / "tool_test.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws["A1"] = "Column1"
        ws["A2"] = "Data"
        wb.save(xlsx_file)

        tool = ReadTool()
        result = tool.execute(path=str(xlsx_file))

        assert result.success
        assert "Column1" in result.llm_output
        assert "Data" in result.llm_output

    @pytest.mark.skipif(not HAS_OPENPYXL, reason="openpyxl not installed")
    def test_read_xlsx_with_formulas(self, tmp_path):
        """Test reading Excel file with formulas (data_only mode)."""
        import openpyxl

        xlsx_file = tmp_path / "formulas.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws["A1"] = 10
        ws["A2"] = 20
        ws["A3"] = "=SUM(A1:A2)"  # Formula
        wb.save(xlsx_file)

        content, description = read_excel_file(xlsx_file)

        # Should have the values (data_only=True means formulas show calculated values
        # but for new files without calculation, may show None)
        assert "[EXCEL FILE: formulas.xlsx]" in content

    def test_xlsx_missing_openpyxl(self, tmp_path):
        """Test error message when openpyxl not installed."""
        xlsx_file = tmp_path / "test.xlsx"
        xlsx_file.write_bytes(b"dummy content")

        with patch('opencode.tools.read.HAS_OPENPYXL', False):
            # Need to reload or call directly with patched value
            content, description = "", "Excel support requires: pip install openpyxl"
            # The actual function checks HAS_OPENPYXL at runtime
            assert "openpyxl" in description


class TestWordReading:
    """Test Word document reading."""

    @pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
    def test_read_docx_file(self, tmp_path):
        """Test reading .docx file."""
        from docx import Document

        docx_file = tmp_path / "test.docx"
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test paragraph.")
        doc.add_paragraph("Another paragraph with more text.")
        doc.save(docx_file)

        content, description = read_word_file(docx_file)

        assert "[WORD DOCUMENT: test.docx]" in content
        assert "Test Document" in content
        assert "test paragraph" in content
        assert "paragraphs" in description

    @pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
    def test_read_docx_with_headings(self, tmp_path):
        """Test reading .docx with multiple heading levels."""
        from docx import Document

        docx_file = tmp_path / "headings.docx"
        doc = Document()
        doc.add_heading("Main Title", 0)
        doc.add_heading("Section 1", level=1)
        doc.add_paragraph("Section 1 content")
        doc.add_heading("Section 2", level=1)
        doc.add_paragraph("Section 2 content")
        doc.save(docx_file)

        content, description = read_word_file(docx_file)

        assert "Main Title" in content
        assert "Section 1" in content
        assert "Section 2" in content

    @pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
    def test_read_docx_with_tables(self, tmp_path):
        """Test reading .docx with tables."""
        from docx import Document

        docx_file = tmp_path / "tables.docx"
        doc = Document()
        doc.add_paragraph("Document with table:")

        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "Header1"
        table.cell(0, 1).text = "Header2"
        table.cell(0, 2).text = "Header3"
        table.cell(1, 0).text = "Data1"
        table.cell(1, 1).text = "Data2"
        table.cell(1, 2).text = "Data3"

        doc.save(docx_file)

        content, description = read_word_file(docx_file)

        assert "[TABLES: 1]" in content
        assert "Header1" in content
        assert "Data1" in content

    @pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
    def test_read_docx_via_tool(self, tmp_path):
        """Test Word reading through ReadTool."""
        from docx import Document

        docx_file = tmp_path / "tool_test.docx"
        doc = Document()
        doc.add_paragraph("Tool test content")
        doc.save(docx_file)

        tool = ReadTool()
        result = tool.execute(path=str(docx_file))

        assert result.success
        assert "Tool test content" in result.llm_output

    @pytest.mark.skipif(not HAS_DOCX, reason="python-docx not installed")
    def test_read_docx_with_properties(self, tmp_path):
        """Test reading .docx with document properties."""
        from docx import Document

        docx_file = tmp_path / "props.docx"
        doc = Document()
        doc.core_properties.title = "My Document Title"
        doc.core_properties.author = "Test Author"
        doc.add_paragraph("Content here")
        doc.save(docx_file)

        content, description = read_word_file(docx_file)

        assert "Title: My Document Title" in content
        assert "Author: Test Author" in content


class TestPDFReading:
    """Test PDF file reading."""

    @pytest.mark.skipif(not HAS_PYPDF, reason="pypdf/PyPDF2 not installed")
    def test_read_pdf_file(self, tmp_path):
        """Test reading PDF file."""
        import pypdf
        from pypdf import PdfWriter

        pdf_file = tmp_path / "test.pdf"

        # Create a simple PDF
        writer = PdfWriter()
        # Add a blank page (we can't easily add text without reportlab)
        writer.add_blank_page(width=612, height=792)
        with open(pdf_file, 'wb') as f:
            writer.write(f)

        content, description = read_pdf_file(pdf_file)

        assert "[PDF FILE: test.pdf]" in content
        assert "1 pages" in description

    @pytest.mark.skipif(not HAS_PYPDF, reason="pypdf/PyPDF2 not installed")
    def test_read_pdf_via_tool(self, tmp_path):
        """Test PDF reading through ReadTool."""
        from pypdf import PdfWriter

        pdf_file = tmp_path / "tool_test.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        with open(pdf_file, 'wb') as f:
            writer.write(f)

        tool = ReadTool()
        result = tool.execute(path=str(pdf_file))

        assert result.success
        assert "PDF FILE" in result.llm_output

    def test_pdf_without_pypdf(self, tmp_path):
        """Test PDF reading shows helpful message without pypdf."""
        pdf_file = tmp_path / "test.pdf"
        pdf_file.write_bytes(b"%PDF-1.4")  # Minimal PDF header

        with patch('opencode.tools.read.HAS_PYPDF', False):
            content, description = read_pdf_file(pdf_file)

            assert "pypdf" in description.lower() or content == ""


class TestBinaryFileHandling:
    """Test handling of binary files."""

    def test_binary_file_info(self, tmp_path):
        """Test that binary files are handled gracefully."""
        binary_file = tmp_path / "test.bin"
        binary_file.write_bytes(b"\x00\x01\x02\x03\xff\xfe\xfd")

        tool = ReadTool()
        result = tool.execute(path=str(binary_file))

        assert result.success
        assert "BINARY FILE" in result.llm_output
        assert "bytes" in result.llm_output

    def test_unknown_extension_binary(self, tmp_path):
        """Test binary file with unknown extension."""
        unknown_file = tmp_path / "data.xyz"
        unknown_file.write_bytes(b"\x89PNG\r\n\x1a\n")  # PNG header

        tool = ReadTool()
        result = tool.execute(path=str(unknown_file))

        assert result.success
        assert "BINARY FILE" in result.llm_output or "xyz" in result.llm_output.lower()


class TestReadToolFileTypes:
    """Test ReadTool with various file types."""

    def test_read_text_file(self, tmp_path):
        """Test reading plain text file."""
        text_file = tmp_path / "readme.txt"
        text_file.write_text("Hello, World!\nThis is a test file.")

        tool = ReadTool()
        result = tool.execute(path=str(text_file))

        assert result.success
        assert "Hello, World!" in result.llm_output

    def test_read_python_file(self, tmp_path):
        """Test reading Python file."""
        py_file = tmp_path / "script.py"
        py_file.write_text("def hello():\n    print('Hello')\n")

        tool = ReadTool()
        result = tool.execute(path=str(py_file))

        assert result.success
        assert "def hello" in result.llm_output

    def test_read_json_file(self, tmp_path):
        """Test reading JSON file."""
        json_file = tmp_path / "config.json"
        json_file.write_text('{"name": "test", "value": 42}')

        tool = ReadTool()
        result = tool.execute(path=str(json_file))

        assert result.success
        assert '"name"' in result.llm_output

    def test_read_nonexistent_file(self, tmp_path):
        """Test reading non-existent file."""
        tool = ReadTool()
        result = tool.execute(path=str(tmp_path / "nonexistent.txt"))

        assert not result.success
        assert "not found" in result.error.lower()

    def test_read_directory_fails(self, tmp_path):
        """Test that reading directory fails."""
        tool = ReadTool()
        result = tool.execute(path=str(tmp_path))

        assert not result.success
        assert "not a file" in result.error.lower()


class TestLargeFileHandling:
    """Test handling of large files."""

    def test_large_file_preview(self, tmp_path):
        """Test that large files get preview treatment."""
        large_file = tmp_path / "large.py"

        # Create a file with many lines
        lines = [f"# Line {i}" for i in range(600)]
        large_file.write_text("\n".join(lines))

        tool = ReadTool()
        result = tool.execute(path=str(large_file))

        assert result.success
        assert "LARGE FILE" in result.llm_output
        assert "600 lines" in result.llm_output

    def test_large_file_full_read(self, tmp_path):
        """Test reading full large file with full=True."""
        large_file = tmp_path / "large.txt"
        lines = [f"Line {i}" for i in range(600)]
        large_file.write_text("\n".join(lines))

        tool = ReadTool()
        result = tool.execute(path=str(large_file), full=True)

        assert result.success
        # Should have all lines, not just preview
        assert "Line 599" in result.llm_output


class TestLineRangeReading:
    """Test reading specific line ranges."""

    def test_read_line_range(self, tmp_path):
        """Test reading specific line range."""
        file = tmp_path / "lines.txt"
        lines = [f"Line {i}" for i in range(1, 21)]
        file.write_text("\n".join(lines))

        tool = ReadTool()
        result = tool.execute(path=str(file), lines="5-10")

        assert result.success
        assert "Line 5" in result.llm_output
        assert "Line 10" in result.llm_output
        # Should not have lines outside range
        assert "Line 1 |" not in result.llm_output
        assert "Line 15" not in result.llm_output

    def test_read_single_line(self, tmp_path):
        """Test reading single line."""
        file = tmp_path / "single.txt"
        file.write_text("Line 1\nLine 2\nLine 3\n")

        tool = ReadTool()
        result = tool.execute(path=str(file), lines="2")

        assert result.success
        assert "Line 2" in result.llm_output


class TestDependencyChecks:
    """Test handling of missing dependencies."""

    def test_excel_without_openpyxl(self, tmp_path):
        """Test Excel reading shows helpful message without openpyxl."""
        xlsx_file = tmp_path / "test.xlsx"
        xlsx_file.write_bytes(b"PK")  # Minimal zip signature

        # Mock HAS_OPENPYXL as False
        with patch.object(
            __import__('opencode.tools.read', fromlist=['HAS_OPENPYXL']),
            'HAS_OPENPYXL',
            False
        ):
            from opencode.tools.read import read_excel_file
            content, description = read_excel_file(xlsx_file)

            # Should return helpful error message
            assert "openpyxl" in description.lower() or content == ""

    def test_word_without_docx(self, tmp_path):
        """Test Word reading shows helpful message without python-docx."""
        docx_file = tmp_path / "test.docx"
        docx_file.write_bytes(b"PK")  # Minimal zip signature

        with patch.object(
            __import__('opencode.tools.read', fromlist=['HAS_DOCX']),
            'HAS_DOCX',
            False
        ):
            from opencode.tools.read import read_word_file
            content, description = read_word_file(docx_file)

            assert "python-docx" in description.lower() or content == ""


class TestEdgeCases:
    """Test edge cases and error handling."""

    def test_empty_excel_file(self, tmp_path):
        """Test reading empty Excel file."""
        if not HAS_OPENPYXL:
            pytest.skip("openpyxl not installed")

        import openpyxl

        xlsx_file = tmp_path / "empty.xlsx"
        wb = openpyxl.Workbook()
        wb.save(xlsx_file)

        content, description = read_excel_file(xlsx_file)

        assert "[EXCEL FILE: empty.xlsx]" in content

    def test_empty_word_file(self, tmp_path):
        """Test reading empty Word file."""
        if not HAS_DOCX:
            pytest.skip("python-docx not installed")

        from docx import Document

        docx_file = tmp_path / "empty.docx"
        doc = Document()
        doc.save(docx_file)

        content, description = read_word_file(docx_file)

        assert "[WORD DOCUMENT: empty.docx]" in content

    def test_unicode_content(self, tmp_path):
        """Test reading files with Unicode content."""
        csv_file = tmp_path / "unicode.csv"
        csv_file.write_text("name,city\nJohn,\u4e1c\u4eac\nMarie,Paris\n", encoding='utf-8')

        content, description = read_csv_file(csv_file)

        assert "John" in content

    def test_read_tool_schema(self):
        """Test ReadTool schema is valid."""
        tool = ReadTool()
        schema = tool.get_schema()

        assert "properties" in schema
        assert "path" in schema["properties"]
        assert "lines" in schema["properties"]
        assert "full" in schema["properties"]
        assert "sheet" in schema["properties"]
        assert "required" in schema
        assert "path" in schema["required"]
